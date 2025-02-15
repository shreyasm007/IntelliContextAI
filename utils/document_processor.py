import PyPDF2
import docx
from io import BytesIO
from typing import List, Tuple, Dict
import numpy as np
from numpy.linalg import norm
import streamlit as st
from .document_preprocessor import DocumentPreprocessor

class MinHashEncoder:
    def __init__(self, num_permutations=128):
        self.num_permutations = num_permutations
        self.prime = 2**31 - 1
        self.permutations = [
            (np.random.randint(1, self.prime), np.random.randint(0, self.prime))
            for _ in range(num_permutations)
        ]

    def _hash_function(self, token: str, a: int, b: int) -> int:
        """Simple string hashing function."""
        hash_val = sum(ord(c) * (31 ** i) for i, c in enumerate(token))
        return ((a * hash_val + b) % self.prime)

    def encode_text(self, text: str) -> np.ndarray:
        """Encode text using MinHash technique."""
        tokens = set(text.lower().split())
        signature = np.array([
            min(self._hash_function(token, a, b) for token in tokens)
            for a, b in self.permutations
        ], dtype=np.float32)
        # Normalize the signature
        if norm(signature) > 0:
            signature = signature / norm(signature)
        return signature

class Document:
    """Document class with enhanced metadata support."""
    def __init__(self, page_content: str, metadata: Dict = None):
        self.page_content = page_content
        self.metadata = metadata or {}

class VectorStore:
    """Custom vector store implementation."""
    def __init__(self, embeddings: MinHashEncoder):
        self.embeddings = embeddings
        self.documents: List[Document] = []
        self.vectors: List[np.ndarray] = []

    def add_documents(self, documents: List[Document]) -> None:
        """Add documents to the vector store."""
        for doc in documents:
            embedding = self.embeddings.encode_text(doc.page_content)
            self.documents.append(doc)
            self.vectors.append(embedding)

    def similarity_search(self, query: str, k: int = 3) -> List[Document]:
        """Perform similarity search using cosine similarity."""
        if not self.vectors:
            return []

        query_embedding = self.embeddings.encode_text(query)
        similarities = np.array([
            np.dot(query_embedding, doc_embedding)
            for doc_embedding in self.vectors
        ])

        # Get top k most similar documents
        top_k_indices = np.argsort(similarities)[-k:][::-1]
        return [self.documents[i] for i in top_k_indices]

class DocumentProcessor:
    def __init__(self):
        self.preprocessor = DocumentPreprocessor()
        self.chunk_size = 500
        self.chunk_overlap = 50
        self.embeddings = MinHashEncoder(num_permutations=64)

    def process_file(self, file_bytes: bytes, file_type: str, filename: str) -> str:
        """Process different file types with enhanced preprocessing."""
        try:
            # Extract raw text based on file type
            if file_type == 'pdf':
                text = self.process_pdf(file_bytes)
            elif file_type == 'docx':
                text = self.process_docx(file_bytes)
            elif file_type == 'txt':
                text = file_bytes.decode('utf-8')
            else:
                raise ValueError(f"Unsupported file type: {file_type}")

            # Validate content
            is_valid, validation_message = self.preprocessor.validate_content(text)
            if not is_valid:
                raise ValueError(validation_message)
            elif validation_message:
                st.warning(validation_message)

            # Clean and preprocess text
            cleaned_text = self.preprocessor.clean_text(text)

            # Extract metadata
            metadata = self.preprocessor.extract_metadata(cleaned_text, filename)
            st.info(f"""📄 Document Stats:
            - Words: {metadata['word_count']:,}
            - Reading Time: {metadata['estimated_reading_time']:.1f} minutes""")

            return cleaned_text
        except Exception as e:
            raise RuntimeError(f"Failed to process file: {str(e)}")

    def process_pdf(self, pdf_file: bytes) -> str:
        """Process PDF file with enhanced error handling."""
        try:
            pdf_reader = PyPDF2.PdfReader(BytesIO(pdf_file))
            text = ""
            total_pages = len(pdf_reader.pages)

            # Show progress bar for PDF processing
            progress_bar = st.progress(0)
            for i, page in enumerate(pdf_reader.pages):
                text += page.extract_text() + "\n"
                progress_bar.progress((i + 1) / total_pages)
            progress_bar.empty()

            return text
        except Exception as e:
            raise RuntimeError(f"Failed to process PDF: {str(e)}")

    def process_docx(self, docx_file: bytes) -> str:
        """Process DOCX file with enhanced error handling."""
        try:
            doc = docx.Document(BytesIO(docx_file))
            text = ""
            total_paragraphs = len(doc.paragraphs)

            # Show progress bar for DOCX processing
            progress_bar = st.progress(0)
            for i, paragraph in enumerate(doc.paragraphs):
                text += paragraph.text + "\n"
                progress_bar.progress((i + 1) / total_paragraphs)
            progress_bar.empty()

            return text
        except Exception as e:
            raise RuntimeError(f"Failed to process DOCX: {str(e)}")

    def create_documents(self, text: str, metadata: Dict = None) -> List[Document]:
        """Create documents with improved chunking and metadata."""
        words = text.split()
        chunks = []

        # Show progress bar for chunking
        progress_bar = st.progress(0)
        chunk_count = max(1, len(words) // (self.chunk_size - self.chunk_overlap))

        i = 0
        while i < len(words):
            chunk = " ".join(words[i:i + self.chunk_size])
            chunk_metadata = metadata.copy() if metadata else {}
            chunk_metadata.update({
                'chunk_index': len(chunks),
                'chunk_size': len(chunk.split())
            })
            chunks.append(Document(chunk, chunk_metadata))
            i += self.chunk_size - self.chunk_overlap
            progress_bar.progress(min(1.0, len(chunks) / chunk_count))

        progress_bar.empty()
        return chunks

    def create_embeddings(self, documents: List[Document]) -> VectorStore:
        """Create embeddings with progress tracking."""
        try:
            vectorstore = VectorStore(self.embeddings)

            # Show progress bar for embedding creation
            progress_bar = st.progress(0)
            total_docs = len(documents)

            for i, doc in enumerate(documents):
                vectorstore.add_documents([doc])
                progress_bar.progress((i + 1) / total_docs)

            progress_bar.empty()
            return vectorstore
        except Exception as e:
            raise RuntimeError(f"Failed to create embeddings: {str(e)}")